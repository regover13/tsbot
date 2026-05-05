#!/usr/bin/env python3
"""
Protokoll-Generator mit Claude API
- Transkript-Segmente werden per KI den Agenda-Punkten zugeordnet
- Teilnehmer werden per Claude Vision aus allen PNGs im Ordner extrahiert
  (oder direkt als Liste übergeben, z.B. aus ServerQuery im Server-Modus)
"""

import sys
import os
import re
import json
import base64
from datetime import datetime, timedelta
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# ── Deutsche ICAO-Codes (Quelle: Wikipedia, vollständig) ──────
ICAO_DE = {
    "EDAB": "Bautzen", "EDAC": "Altenburg", "EDAD": "Dessau", "EDAE": "Eisenhüttenstadt",
    "EDAG": "Großrückerswalde", "EDAH": "Heringsdorf", "EDAI": "Segeletz", "EDAJ": "Gera",
    "EDAK": "Großenhain", "EDAL": "Fürstenwalde", "EDAM": "Merseburg", "EDAN": "Neustadt-Glewe",
    "EDAO": "Nordhausen", "EDAP": "Neuhausen", "EDAQ": "Halle", "EDAR": "Pirna",
    "EDAS": "Finsterwalde", "EDAT": "Nardt", "EDAU": "Riesa", "EDAV": "Eberswalde",
    "EDAW": "Roitzschjora", "EDAX": "Müritz", "EDAY": "Strausberg", "EDAZ": "Schönhagen",
    "EDBA": "Arnstadt", "EDBC": "Magdeburg", "EDBD": "Dedelow", "EDBE": "Brandenburg",
    "EDBF": "Fehrbellin", "EDBG": "Burg", "EDBH": "Barth", "EDBI": "Zwickau",
    "EDBJ": "Jena", "EDBK": "Kyritz", "EDBL": "Laucha", "EDBM": "Magdeburg",
    "EDBN": "Neubrandenburg", "EDBO": "Oehna", "EDBP": "Pinnow", "EDBQ": "Bronkow",
    "EDBR": "Rothenburg", "EDBS": "Sömmerda", "EDBT": "Allstedt", "EDBU": "Pritzwalk",
    "EDBV": "Stralsund", "EDBW": "Werneuchen", "EDBX": "Görlitz", "EDBY": "Schmoldow",
    "EDBZ": "Schwarzheide", "EDCA": "Anklam", "EDCB": "Ballenstedt", "EDCD": "Cottbus",
    "EDCE": "Eggersdorf", "EDCF": "Friedersdorf", "EDCG": "Rügen", "EDCH": "Sprossen",
    "EDCI": "Klix", "EDCJ": "Chemnitz", "EDCK": "Köthen", "EDCL": "Klietz",
    "EDCM": "Kamenz", "EDCO": "Obermehler", "EDCP": "Peenemünde", "EDCQ": "Aschersleben",
    "EDCR": "Rerik", "EDCS": "Saarmund", "EDCT": "Taucha", "EDCU": "Güstrow",
    "EDCV": "Pasewalk", "EDCW": "Wismar", "EDCX": "Purkshof", "EDCY": "Welzow",
    "EDDB": "Berlin (BER)", "EDDC": "Dresden", "EDDE": "Erfurt", "EDDF": "Frankfurt",
    "EDDG": "Münster/Osnabrück", "EDDH": "Hamburg", "EDDI": "Berlin (Tempelhof)",
    "EDDK": "Köln/Bonn", "EDDL": "Düsseldorf", "EDDM": "München", "EDDN": "Nürnberg",
    "EDDP": "Leipzig/Halle", "EDDR": "Saarbrücken", "EDDS": "Stuttgart",
    "EDDT": "Berlin (Tegel)", "EDDV": "Hannover", "EDDW": "Bremen",
    "EDEB": "Bad Langensalza", "EDEF": "Babenhausen", "EDEG": "Gotha", "EDEH": "Herrenteich",
    "EDEK": "Baumholder", "EDEL": "Langenlonsheim", "EDEM": "Mosenberg", "EDEN": "Bad Hersfeld",
    "EDEP": "Heppenheim", "EDEQ": "Mühlhausen", "EDER": "Wasserkuppe", "EDEW": "Walldürn",
    "EDFA": "Anspach", "EDFB": "Reichelsheim", "EDFC": "Aschaffenburg", "EDFD": "Bad Neustadt",
    "EDFE": "Egelsbach", "EDFG": "Gelnhausen", "EDFH": "Hahn", "EDFI": "Hirzenhain",
    "EDFJ": "Hammelburg", "EDFK": "Bad Kissingen", "EDFL": "Gießen", "EDFM": "Mannheim",
    "EDFN": "Marburg", "EDFO": "Michelstadt", "EDFP": "Ober-Mörlen", "EDFQ": "Allendorf",
    "EDFR": "Rothenburg ob der Tauber", "EDFS": "Schweinfurt", "EDFT": "Lauterbach",
    "EDFU": "Mainbullau", "EDFV": "Worms", "EDFW": "Würzburg", "EDFX": "Hockenheim",
    "EDFY": "Elz", "EDFZ": "Mainz", "EDGA": "Ailertchen", "EDGB": "Breitscheid",
    "EDGE": "Eisenach", "EDGF": "Fulda", "EDGH": "Hettstadt", "EDGI": "Ingelfingen",
    "EDGJ": "Ochsenfurt", "EDGK": "Korbach", "EDGM": "Mosbach", "EDGN": "Nordenbeck",
    "EDGP": "Oppenheim", "EDGQ": "Schameder", "EDGR": "Gießen", "EDGS": "Siegerland",
    "EDGT": "Bottenhorn", "EDGU": "Unterschüpf", "EDGW": "Wolfhagen", "EDGX": "Walldorf",
    "EDGY": "Kitzingen", "EDGZ": "Weinheim", "EDHA": "Ahlhorn", "EDHB": "Grube",
    "EDHC": "Lüchow", "EDHD": "Eichsfeld", "EDHE": "Uetersen", "EDHF": "Itzehoe",
    "EDHG": "Lüneburg", "EDHI": "Hamburg-Finkenwerder", "EDHK": "Kiel", "EDHL": "Lübeck",
    "EDHM": "Hartenholm", "EDHN": "Neumünster", "EDHO": "Ahrenlohe", "EDHP": "Pellworm",
    "EDHS": "Stade", "EDHU": "Lauenbrück", "EDHW": "Wahlstedt", "EDHY": "Hoya",
    "EDIU": "Heidelberg", "EDJA": "Memmingen", "EDJG": "Grabenstätt", "EDJR": "Saarlouis",
    "EDKA": "Aachen", "EDKB": "Bonn", "EDKD": "Altena", "EDKF": "Bergneustadt",
    "EDKH": "Hünsborn", "EDKI": "Betzdorf", "EDKL": "Leverkusen", "EDKM": "Meschede",
    "EDKN": "Wipperfürth", "EDKO": "Brilon", "EDKP": "Plettenberg", "EDKR": "Schmallenberg",
    "EDKU": "Attendorn", "EDKV": "Dahlemer Binz", "EDKW": "Werdohl", "EDKZ": "Meinerzhagen",
    "EDLA": "Arnsberg", "EDLB": "Borkenberge", "EDLC": "Kamp-Lintfort", "EDLD": "Dinslaken",
    "EDLE": "Essen", "EDLF": "Grefrath", "EDLG": "Goch", "EDLH": "Hamm",
    "EDLI": "Bielefeld", "EDLJ": "Detmold", "EDLK": "Krefeld", "EDLM": "Marl",
    "EDLN": "Mönchengladbach", "EDLO": "Oerlinghausen", "EDLP": "Paderborn/Lippstadt",
    "EDLQ": "Beelen", "EDLR": "Paderborn", "EDLS": "Stadtlohn", "EDLT": "Münster",
    "EDLU": "Oelde", "EDLV": "Wesel", "EDLW": "Dortmund", "EDLX": "Wesel",
    "EDLY": "Borken", "EDLZ": "Soest", "EDMA": "Augsburg", "EDMB": "Biberach",
    "EDMC": "Blaubeuren", "EDMD": "Dachau", "EDME": "Eggenfelden", "EDMF": "Fürstenzell",
    "EDMG": "Günzburg", "EDMH": "Gunzenhausen", "EDMI": "Illertissen", "EDMJ": "Jesenwang",
    "EDMK": "Kempten", "EDML": "Landshut", "EDMN": "Mindelheim", "EDMO": "Oberpfaffenhofen",
    "EDMP": "Vilsbiburg", "EDMQ": "Donauwörth", "EDMS": "Straubing", "EDMT": "Tannheim",
    "EDMU": "Gundelfingen", "EDMV": "Vilshofen", "EDMW": "Deggendorf", "EDMY": "Mühldorf",
    "EDMZ": "Bad Waldsee", "EDNA": "Ampfing", "EDNB": "Arnbruck", "EDNC": "Beilngries",
    "EDND": "Dinkelsbühl", "EDNE": "Erbach", "EDNF": "Grafenau", "EDNG": "Giengen",
    "EDNH": "Bad Wörishofen", "EDNI": "Berching", "EDNJ": "Neuburg an der Donau",
    "EDNK": "Kirchdorf an der Iller", "EDNL": "Leutkirch", "EDNM": "Nittenau",
    "EDNO": "Nördlingen", "EDNP": "Pfarrkirchen", "EDNQ": "Bopfingen", "EDNR": "Regensburg",
    "EDNS": "Schwabmünchen", "EDNT": "Treuchtlingen", "EDNU": "Thannhausen",
    "EDNV": "Vogtareuth", "EDNW": "Weißenhorn", "EDNX": "Schleißheim",
    "EDNY": "Friedrichshafen", "EDNZ": "Zell am Harmersbach",
    "EDOA": "Auerbach", "EDOB": "Bad Berka", "EDOC": "Gardelegen", "EDOD": "Reinsdorf",
    "EDOE": "Schwaighofen", "EDOF": "Bad Frankenhausen", "EDOG": "Torgau",
    "EDOH": "Langhennersdorf", "EDOI": "Bienenfarm", "EDOJ": "Lüsse", "EDOK": "Rudolstadt",
    "EDOL": "Oschersleben", "EDOM": "Klein Mühlingen", "EDON": "Neuhardenberg",
    "EDOP": "Schwerin", "EDOQ": "Oschatz", "EDOR": "Stölln", "EDOS": "Pennewitz",
    "EDOT": "Greiz", "EDOU": "Weimar", "EDOV": "Stendal", "EDOW": "Waren",
    "EDOX": "Renneritz", "EDOZ": "Schönebeck",
    "EDPA": "Aalen", "EDPB": "Bad Ditzenbach", "EDPC": "Bad Endorf", "EDPD": "Dingolfing",
    "EDPE": "Eichstätt", "EDPF": "Schwandorf", "EDPG": "Griesau", "EDPH": "Neuhausen ob Eck",
    "EDPI": "Moosburg", "EDPJ": "Laichingen", "EDPK": "Schönberg", "EDPM": "Donzdorf",
    "EDPO": "Neumarkt in der Oberpfalz", "EDPQ": "Schmidgaden", "EDPS": "Sonnen",
    "EDPT": "Gerstetten", "EDPU": "Bartholomä", "EDPW": "Thalmässing", "EDPY": "Ellwangen",
    "EDQA": "Bamberg", "EDQB": "Bad Windsheim", "EDQC": "Coburg", "EDQD": "Bayreuth",
    "EDQE": "Burg Feuerstein", "EDQF": "Ansbach", "EDQG": "Giebelstadt",
    "EDQH": "Herzogenaurach", "EDQI": "Lauf an der Pegnitz", "EDQK": "Kulmbach",
    "EDQL": "Lichtenfels", "EDQM": "Hof", "EDQN": "Neustadt an der Aisch",
    "EDQO": "Ottengrüner Heide", "EDQP": "Rosenthal", "EDQR": "Ebern", "EDQS": "Suhl",
    "EDQT": "Haßfurt", "EDQW": "Weiden", "EDQX": "Hetzleser Berg", "EDQY": "Coburg",
    "EDQZ": "Pegnitz",
    "EDRA": "Bad Neuenahr", "EDRB": "Bitburg", "EDRD": "Neumagen-Dhron", "EDRE": "Mendig",
    "EDRF": "Bad Dürkheim", "EDRG": "Idar-Oberstein", "EDRH": "Hoppstädten-Weiersbach",
    "EDRI": "Linkenheim", "EDRJ": "Saarlouis", "EDRK": "Koblenz", "EDRL": "Lachen-Speyerdorf",
    "EDRM": "Traben-Trarbach", "EDRN": "Nannhausen", "EDRO": "Schweighofen",
    "EDRP": "Pirmasens", "EDRS": "Bad Sobernheim", "EDRT": "Trier",
    "EDRV": "Wershofen", "EDRW": "Dierdorf", "EDRX": "Neunkirchen", "EDRY": "Speyer",
    "EDRZ": "Zweibrücken",
    "EDSA": "Albstadt-Degerfeld", "EDSB": "Karlsruhe/Baden-Baden", "EDSD": "Deckenpfronn",
    "EDSE": "Göppingen", "EDSF": "Hütten", "EDSG": "Grabenstetten", "EDSH": "Backnang",
    "EDSI": "Binningen", "EDSK": "Kehl", "EDSL": "Blumberg", "EDSM": "Müllheim",
    "EDSN": "Neuhausen ob Eck", "EDSO": "Gruibingen", "EDSP": "Pferdsfeld",
    "EDSR": "Radolfzell", "EDST": "Kirchheim/Hahnweide", "EDSW": "Altdorf",
    "EDSX": "Völkleshofen", "EDSZ": "Rottweil",
    "EDTA": "Bohlhof", "EDTB": "Baden-Oos", "EDTC": "Bruchsal", "EDTD": "Donaueschingen",
    "EDTE": "Eutingen im Gäu", "EDTF": "Freiburg im Breisgau", "EDTG": "Bremgarten",
    "EDTH": "Heubach", "EDTK": "Karlsruhe", "EDTL": "Lahr", "EDTM": "Mengen",
    "EDTN": "Nabern", "EDTO": "Offenburg", "EDTP": "Pfullendorf", "EDTQ": "Pattonville",
    "EDTR": "Rheinfelden", "EDTS": "Schwenningen", "EDTU": "Saulgau", "EDTW": "Schramberg",
    "EDTX": "Schwäbisch Hall", "EDTY": "Schwäbisch Hall", "EDTZ": "Konstanz",
    "EDUA": "Stechow", "EDUB": "Brandenburg an der Havel", "EDUF": "Falkenberg",
    "EDUG": "Gransee", "EDUO": "Oberrißdorf", "EDUP": "Perleberg", "EDUR": "Brüggen",
    "EDUS": "Finsterwalde", "EDUT": "Templin", "EDUW": "Tutow", "EDUY": "Sedlitzer See",
    "EDUZ": "Zerbst",
    "EDVA": "Bad Gandersheim", "EDVC": "Celle", "EDVD": "Uslar", "EDVE": "Braunschweig",
    "EDVF": "Blomberg", "EDVG": "Mengeringhausen", "EDVH": "Hodenhagen", "EDVI": "Höxter",
    "EDVJ": "Salzgitter", "EDVK": "Kassel", "EDVL": "Hölleberg", "EDVM": "Hildesheim",
    "EDVN": "Northeim", "EDVP": "Peine", "EDVQ": "Wilsche", "EDVR": "Rinteln",
    "EDVS": "Salzgitter", "EDVU": "Uelzen", "EDVW": "Bad Pyrmont", "EDVY": "Porta Westfalica",
    "EDWA": "Bordelum", "EDWB": "Bremerhaven", "EDWC": "Damme", "EDWD": "Lemwerder",
    "EDWE": "Emden", "EDWF": "Leer", "EDWG": "Wangerooge", "EDWH": "Oldenburg",
    "EDWI": "Wilhelmshaven", "EDWJ": "Juist", "EDWK": "Karlshöfen", "EDWL": "Langeoog",
    "EDWM": "Weser-Wümme", "EDWN": "Nordhorn", "EDWO": "Osnabrück", "EDWP": "Wiefelstede",
    "EDWQ": "Ganderkesee", "EDWR": "Borkum", "EDWS": "Norden", "EDWT": "Blexen",
    "EDWU": "Varrelbusch", "EDWV": "Verden", "EDWX": "Westerstede", "EDWY": "Norderney",
    "EDWZ": "Baltrum",
    "EDXA": "Achmer", "EDXB": "Heide-Büsum", "EDXC": "Schleswig", "EDXD": "Bohmte",
    "EDXE": "Rheine", "EDXF": "Flensburg", "EDXG": "Melle", "EDXH": "Helgoland",
    "EDXI": "Nienburg", "EDXJ": "Husum", "EDXK": "Leck", "EDXL": "Barßel",
    "EDXM": "St. Michaelisdonn", "EDXN": "Nordholz", "EDXO": "St. Peter-Ording",
    "EDXP": "Harle", "EDXQ": "Rotenburg", "EDXR": "Rendsburg", "EDXS": "Seedorf",
    "EDXT": "Sierksdorf", "EDXU": "Hüttenbusch", "EDXW": "Sylt", "EDXY": "Wyk auf Föhr",
    "EDXZ": "Kührstedt",
    "ETAD": "Spangdahlem", "ETAR": "Ramstein", "ETBS": "Berlin", "ETEB": "Ansbach",
    "ETED": "Kaiserslautern", "ETEJ": "Bamberg", "ETEK": "Baumholder", "ETGU": "Ulm",
    "ETHA": "Altenstadt", "ETHB": "Bückeburg", "ETHC": "Celle", "ETHE": "Rheine",
    "ETHF": "Fritzlar", "ETHI": "Itzehoe", "ETHL": "Laupheim", "ETHM": "Mendig",
    "ETHN": "Niederstetten", "ETHR": "Roth", "ETHS": "Faßberg", "ETIC": "Grafenwöhr",
    "ETID": "Langendiebach", "ETIE": "Heidelberg", "ETIH": "Hohenfels", "ETIK": "Illesheim",
    "ETIN": "Kitzingen", "ETLS": "Leipzig", "ETME": "Eggebek", "ETMK": "Kiel",
    "ETMN": "Nordholz", "ETND": "Diepholz", "ETNG": "Geilenkirchen", "ETNH": "Hohn",
    "ETNJ": "Jever", "ETNL": "Rostock", "ETNN": "Nörvenich", "ETNP": "Hopsten",
    "ETNS": "Schleswig", "ETNT": "Wittmund", "ETNU": "Neubrandenburg", "ETNW": "Wunstorf",
    "ETOI": "Vilseck", "ETOR": "Coleman Barracks", "ETOU": "Wiesbaden", "ETOY": "Leighton",
    "ETSA": "Landsberg am Lech", "ETSB": "Büchel", "ETSE": "Erding",
    "ETSF": "Fürstenfeldbruck", "ETSH": "Holzdorf", "ETSI": "Ingolstadt", "ETSL": "Lechfeld",
    "ETSN": "Neuburg an der Donau", "ETSP": "Pferdsfeld", "ETSR": "Roth",
    "ETUO": "Gütersloh", "ETUR": "Brüggen", "ETWM": "Meppen",
}


def _icao_block() -> str:
    lines = [f"{k} → {v}" for k, v in sorted(ICAO_DE.items())]
    return (
        "ICAO-FLUGHAFENREFERENZ (verbindlich — nur diese Zuordnungen verwenden):\n"
        + "\n".join(lines)
        + "\nFür ICAO-Codes die oben nicht aufgeführt sind: Code unverändert übernehmen, nicht raten."
    )


# ── Konfiguration ─────────────────────────────────────────────
def lese_config(skript_ordner: str) -> dict:
    # Env-Variablen haben Vorrang; config.txt überschreibt nur wenn vorhanden
    config = {
        "ANTHROPIC_API_KEY": os.environ.get("ANTHROPIC_API_KEY"),
        "CLAUDE_MODEL": os.environ.get("CLAUDE_MODEL", "claude-sonnet-4-5-20250929"),
    }
    pfad = os.path.join(skript_ordner, "config.txt")
    if os.path.exists(pfad):
        with open(pfad, "r", encoding="utf-8") as f:
            for zeile in f:
                zeile = zeile.strip()
                if zeile.startswith("#") or "=" not in zeile:
                    continue
                key, _, wert = zeile.partition("=")
                config[key.strip()] = wert.strip()
    return config


# ── Hilfsfunktionen ───────────────────────────────────────────
def setze_seitenraender(doc, oben=2.5, unten=2.5, links=3.0, rechts=2.5):
    for section in doc.sections:
        section.top_margin    = Cm(oben)
        section.bottom_margin = Cm(unten)
        section.left_margin   = Cm(links)
        section.right_margin  = Cm(rechts)


def lese_agenda(pfad: str) -> list:
    if not pfad or not os.path.exists(pfad):
        return []
    with open(pfad, "r", encoding="utf-8") as f:
        return [z.strip() for z in f if z.strip()]


def lese_transkript(pfad: str) -> tuple:
    with open(pfad, "r", encoding="utf-8") as f:
        inhalt = f.read()
    muster = re.compile(r'\[(\d+:\d{2}) - (\d+:\d{2})\] (.+)')
    segmente = muster.findall(inhalt)
    volltext_match = re.search(r'VOLLTEXT:\s*\n\n(.+)', inhalt, re.DOTALL)
    volltext = volltext_match.group(1).strip() if volltext_match else inhalt
    return volltext, segmente


def _mm_ss_zu_uhrzeit(zeitraum: str, basis: datetime) -> str:
    """Konvertiert 'MM:SS - MM:SS' → 'HH:MM - HH:MM Uhr' (absolute Uhrzeit)."""
    pattern = re.compile(r'(\d+):(\d{2})\s*-\s*(\d+):(\d{2})')
    m = pattern.match(zeitraum.strip())
    if not m or not basis:
        return zeitraum
    def to_uhr(mm, ss):
        dt = basis + timedelta(minutes=int(mm), seconds=int(ss))
        return dt.strftime("%H:%M")
    return f"{to_uhr(m.group(1), m.group(2))} - {to_uhr(m.group(3), m.group(4))} Uhr"


def _schneide_transkript(zeilen: list[str], start_mm_ss: str,
                         ende_mm_ss: str | None) -> str:
    """Extrahiert Transkript-Zeilen anhand von MM:SS-Grenzen (+ 2 min Overlap am Anfang)."""
    def zu_sekunden(mm_ss: str) -> int:
        teile = mm_ss.split(":")
        return int(teile[0]) * 60 + int(teile[1])

    start_s = max(0, zu_sekunden(start_mm_ss) - 120)
    ende_s  = zu_sekunden(ende_mm_ss) + 30 if ende_mm_ss else float("inf")

    muster = re.compile(r'\[(\d+):(\d{2})')
    ergebnis = []
    for zeile in zeilen:
        m = muster.search(zeile)
        if m:
            sek = int(m.group(1)) * 60 + int(m.group(2))
            if start_s <= sek < ende_s:
                ergebnis.append(zeile)
        elif ergebnis:
            ergebnis.append(zeile)
    return "\n".join(ergebnis)


def finde_alle_pngs(ordner: str) -> list:
    """Alle PNG-Dateien im Ordner, sortiert nach Änderungszeit."""
    pngs = [
        os.path.join(ordner, f)
        for f in os.listdir(ordner)
        if f.lower().endswith(".png")
    ]
    return sorted(pngs, key=os.path.getmtime)


# ── Claude Vision: Teilnehmer aus Screenshots extrahieren ─────
def extrahiere_teilnehmer(png_pfade: list, api_key: str, modell: str) -> list:
    """Sendet alle PNGs an Claude Vision und extrahiert Teilnehmernamen."""
    try:
        import anthropic
    except ImportError:
        print("HINWEIS: anthropic-Paket nicht installiert.")
        return []

    if not png_pfade:
        return []

    print(f"Extrahiere Teilnehmer aus {len(png_pfade)} Screenshot(s)...")
    client = anthropic.Anthropic(api_key=api_key)

    alle_roh = []  # Liste von dicts vor Deduplizierung

    for pfad in png_pfade:
        with open(pfad, "rb") as f:
            bild_data = base64.standard_b64encode(f.read()).decode("utf-8")

        message = client.messages.create(
            model=modell,
            max_tokens=1024,
            messages=[{
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": "image/png",
                            "data": bild_data,
                        },
                    },
                    {
                        "type": "text",
                        "text": """Das ist ein TeamSpeak-Screenshot einer Flugsimulations-Community.

Extrahiere NUR echte Personennamen der eingeloggten Nutzer.
Echte Nutzernamen haben IMMER das Format: "Vorname Nachname/FRSxxx" oder "Vorname Nachname/FRSxxx (Zusatz)".

WICHTIG - folgendes sind KEINE Nutzernamen und dürfen NICHT extrahiert werden:
- Kanal- oder Raumbezeichnungen (z.B. "Fliegen Weltweit", "AFK", "Lobby", "Training", etc.)
- Kanalgruppen oder Kategorien
- Statusmeldungen

Antworte NUR mit einer JSON-Liste. Wenn keine gültigen Nutzernamen gefunden, antworte mit [].
Format:
[
  {"name": "Vorname Nachname", "frs": "FRS49"},
  {"name": "Anderer Name", "frs": "FRS12"}
]"""
                    }
                ],
            }]
        )

        antwort = message.content[0].text.strip()
        json_match = re.search(r'\[.*\]', antwort, re.DOTALL)
        if json_match:
            try:
                teilnehmer = json.loads(json_match.group())
                alle_roh.extend(teilnehmer)
            except json.JSONDecodeError:
                pass

    # Fuzzy-Deduplizierung: gleiche FRS → behalte längeren Namen
    def aehnlichkeit(a: str, b: str) -> float:
        a, b = a.lower(), b.lower()
        if not a or not b:
            return 0.0
        treffer = sum(c in b for c in a)
        return treffer / max(len(a), len(b))

    eindeutig = []
    for eintrag in alle_roh:
        name = eintrag.get("name", "").strip()
        frs  = eintrag.get("frs", "").strip()
        if not name:
            continue
        if frs:
            if any(e["frs"] == frs for e in eindeutig):
                for e in eindeutig:
                    if e["frs"] == frs and len(name) > len(e["name"]):
                        e["name"] = name
                continue
        if any(aehnlichkeit(name, e["name"]) > 0.85 for e in eindeutig):
            continue
        eindeutig.append({"name": name, "frs": frs})

    return sorted(eindeutig, key=lambda x: x["name"].lower())


# ── Claude API: Transkript den Agenda-Punkten zuordnen ─────────
def ki_zuordnung(volltext: str, segmente: list, agenda: list, api_key: str, modell: str,
                 extra_instruktionen: str = None,
                 kanal_wechsel: list = None,
                 teilnehmer: list = None,
                 session_started_at: datetime = None) -> tuple:
    try:
        import anthropic
    except ImportError:
        print("HINWEIS: anthropic-Paket nicht installiert.")
        return []

    print("Sende Transkript an Claude API zur Zuordnung...")
    transkript_text = "\n".join([f"[{s} - {e}] {t}" for s, e, t in segmente]) if segmente else volltext
    agenda_text = "\n".join([f"{i+1}. {p}" for i, p in enumerate(agenda)])

    # ── Kanalwechsel-Block ────────────────────────────────────
    kanal_block = ""
    if kanal_wechsel:
        kanal_block = "\nKANALWECHSEL WÄHREND DER SITZUNG (Zeitstempel relativ zum Aufnahmestart):\n"
        for evt in kanal_wechsel:
            von  = evt.get("from_channel_name") or f"Kanal {evt.get('from_channel', '?')}"
            nach = evt.get("to_channel_name")   or f"Kanal {evt.get('to_channel', '?')}"
            if session_started_at and evt.get("timestamp"):
                try:
                    evt_time = datetime.fromisoformat(evt["timestamp"])
                    rel_sec  = max(0, int((evt_time - session_started_at).total_seconds()))
                    rel_str  = f"{rel_sec // 60:02d}:{rel_sec % 60:02d}"
                    kanal_block += f"- [{rel_str}] {von} → {nach}\n"
                except Exception:
                    kanal_block += f"- {von} → {nach}\n"
            else:
                kanal_block += f"- {von} → {nach}\n"
        kanal_block += (
            "Füge jeden Kanalwechsel in die Zusammenfassung des Agenda-Punkts ein, "
            "in dessen MM:SS-Zeitraum er fällt (vergleiche den [MM:SS]-Zeitstempel "
            "des Kanalwechsels mit dem zeitraum-Feld des jeweiligen Agenda-Punkts).\n"
        )

    # ── Zusätzliche Instruktionen ─────────────────────────────
    extra_block = ""
    if extra_instruktionen and extra_instruktionen.strip():
        extra_block = f"\nZUSÄTZLICHE INSTRUKTIONEN DES NUTZERS:\n{extra_instruktionen.strip()}\n"

    # ── ICAO-Referenz ──────────────────────────────────────────
    icao_ref = _icao_block()

    # ── Teilnehmer-Block ──────────────────────────────────────
    teilnehmer_block = ""
    if teilnehmer:
        zeilen = []
        for t in teilnehmer:
            name = t.get("name", "")
            frs  = t.get("frs", "")
            zeilen.append(f"- {name}" + (f" ({frs})" if frs else ""))
        teilnehmer_block = "\nTEILNEHMER (exakte Namen und Kennzeichen – unverändert übernehmen):\n" \
                           + "\n".join(zeilen) + "\n"

    # ── Sitzungsdatum-Block ───────────────────────────────────
    datum_block = ""
    if session_started_at:
        wochentage = ["Montag", "Dienstag", "Mittwoch", "Donnerstag", "Freitag", "Samstag", "Sonntag"]
        wochentag = wochentage[session_started_at.weekday()]
        datum_block = (
            f"\nSITZUNGSDATUM: {wochentag}, {session_started_at.strftime('%d.%m.%Y')}, "
            f"Sitzungsstart: {session_started_at.strftime('%H:%M')} Uhr\n"
            "Nutze dieses Datum, um relative Zeitangaben im Transkript (z.B. 'morgen', 'heute', "
            "'nächste Woche') korrekt in absolute Daten aufzulösen.\n"
        )

    prompt = f"""Du bist ein professioneller Protokollschreiber. Analysiere das Transkript und weise jeden Abschnitt dem passenden Agenda-Punkt zu.
{datum_block}{teilnehmer_block}
AGENDA:
{agenda_text}

TRANSKRIPT:
{transkript_text}
{kanal_block}{extra_block}
{icao_ref}

Antworte NUR mit folgendem JSON:
{{
  "agenda_punkte": [
    {{
      "punkt": "Exakter Name des Agenda-Punkts",
      "zusammenfassung": "Kurze einleitende Zusammenfassung (1-2 Sätze: wer hat was präsentiert)",
      "details": ["Einzelner Aufzählungspunkt", "Weiterer Punkt"],
      "beschluesse": ["Beschluss oder Aktionspunkt 1"],
      "zeitraum": "00:00 - 08:30"
    }}
  ]
}}

Hinweise:
- Jeden Agenda-Punkt aufführen, auch ohne Transkript-Treffer
- Zusammenfassung sachlich und neutral, nur 1-2 einleitende Sätze
- details: Aufzählungsliste für Events, Programmpunkte, Termine, Stichpunkte – leer lassen wenn kein Listeninhalt vorhanden
- Beschlüsse = konkrete Entscheidungen oder Aktionspunkte
- WICHTIG: Antworte mit reinem JSON ohne Markdown-Code-Blöcke (kein ```json)
- WICHTIG: Keine wörtlichen Zitate mit Anführungszeichen in den Strings – paraphrasieren statt zitieren
- NAMEN: Wenn ein Sprecher im Transkript von anderen namentlich angesprochen wird oder sich selbst vorstellt, gleiche den Namen mit der Teilnehmerliste ab und verwende immer den vollständigen Namen. Schreibe nur dann 'ein Mitglied', wenn der Sprecher absolut nicht identifizierbar ist."""

    # ── Zweistufiger Modus für lange Transkripte ─────────────────
    if len(transkript_text) > 80_000 and len(agenda) > 1:
        print(f"Transkript lang ({len(transkript_text):,} Zeichen) → zweistufige Verarbeitung.")
        uebergaenge = ki_segment_timestamps(transkript_text, agenda, api_key, modell)
        if uebergaenge and len(uebergaenge) >= len(agenda) - 1:
            zeilen = transkript_text.splitlines()
            alle_punkte = []
            for i, punkt in enumerate(agenda):
                start = uebergaenge[i]["start"] if i < len(uebergaenge) else "00:00"
                ende  = uebergaenge[i + 1]["start"] if i + 1 < len(uebergaenge) else None
                ausschnitt = _schneide_transkript(zeilen, start, ende)
                print(f"  Punkt {i+1} ({punkt}): {start} – {ende or 'Ende'}, "
                      f"{len(ausschnitt):,} Zeichen")
                teil = ki_zuordnung(
                    ausschnitt, [], [punkt], api_key, modell,
                    extra_instruktionen=extra_instruktionen,
                    kanal_wechsel=kanal_wechsel,
                    teilnehmer=teilnehmer,
                    session_started_at=session_started_at,
                )
                if teil:
                    alle_punkte.extend(teil)
                else:
                    alle_punkte.append({
                        "punkt": punkt,
                        "zusammenfassung": "Kein Inhalt ermittelt.",
                        "details": [], "beschluesse": [], "zeitraum": start,
                    })
            return alle_punkte
        print("  Segmentierung unvollständig – Fallback auf Einzel-Pass.")
    # ── Ende zweistufiger Modus ───────────────────────────────────

    client = anthropic.Anthropic(api_key=api_key)
    message = client.messages.create(
        model=modell,
        max_tokens=8192,
        temperature=0.3,
        messages=[{"role": "user", "content": prompt}]
    )

    antwort = message.content[0].text.strip()
    # Markdown-Code-Blöcke entfernen (```json ... ``` oder ``` ... ```)
    antwort = re.sub(r'^```(?:json)?\s*', '', antwort)
    antwort = re.sub(r'\s*```$', '', antwort).strip()
    print(f"Claude Antwort ({len(antwort)} Zeichen, stop_reason={message.stop_reason}): {antwort[:200]}...")
    json_match = re.search(r'\{.*\}', antwort, re.DOTALL)
    if json_match:
        json_str = json_match.group()
        for versuch, text in enumerate([json_str,
                                        re.sub(r'(?<=[^\\])"(?=[^,\]\}:\n])', r'\\"', json_str)]):
            try:
                data = json.loads(text)
                punkte = data.get("agenda_punkte", [])
                print(f"JSON OK (Versuch {versuch+1}) – {len(punkte)} Agenda-Punkte geparst.")
                return punkte
            except (json.JSONDecodeError, KeyError) as e:
                print(f"JSON-Parse-Fehler (Versuch {versuch+1}): {e}")
    else:
        print(f"Kein JSON in Antwort gefunden. Antwort: {antwort[:500]}")
    return []


# ── Segmentierungs-Hilfsfunktion (für zweistufigen Modus) ────────
def ki_segment_timestamps(transkript_text: str, agenda: list,
                          api_key: str, modell: str) -> list[dict]:
    """
    Pass 1 des zweistufigen Modus: Ermittelt via Claude den Startzeit-Stempel
    jedes Agendapunkts. Claude erkennt aus dem Kontext formale Ankündigungen
    und unterscheidet sie von Forward-References ("das klären wir später").
    """
    try:
        import anthropic
    except ImportError:
        return []

    agenda_text = "\n".join(f"{i+1}. {p}" for i, p in enumerate(agenda))
    prompt = (
        "Im folgenden Transkript stehen Zeitstempel im Format [MM:SS - MM:SS].\n"
        "Identifiziere für jeden Agendapunkt den Zeitstempel, ab dem er FORMAL beginnt.\n\n"
        "WICHTIG:\n"
        "- Nur die formale Ankündigung durch den Moderator zählt\n"
        "  (z.B. 'jetzt kommen wir zu...', 'als nächstes...', 'wir wären beim Punkt...')\n"
        "- Eine Nennung eines späteren Punkts WÄHREND eines laufenden Punkts\n"
        "  ('das besprechen wir unter Anliegen der Mitglieder') zählt NICHT als Start\n"
        "- 'Anliegen der Mitglieder' ist ein offener Sammelkorb für Spontanthemen;\n"
        "  er beginnt erst wenn der Moderator ihn offiziell eröffnet\n"
        "- 'Internes (nicht öffentlich)' ähnlich: beginnt erst bei formaler Eröffnung,\n"
        "  typischerweise mit einem Hinweis auf kleinere/nichtöffentliche Runde\n"
        "- Schätze fehlende Punkte anhand der Reihenfolge und des Zeitverlaufs\n\n"
        f"AGENDA:\n{agenda_text}\n\n"
        f"TRANSKRIPT:\n{transkript_text}\n\n"
        "Antworte NUR mit JSON (kein Markdown):\n"
        '{"uebergaenge": [{"nummer": 1, "start": "07:19"}, ...]}'
    )
    client = anthropic.Anthropic(api_key=api_key)
    msg = client.messages.create(
        model=modell, max_tokens=512, temperature=0,
        messages=[{"role": "user", "content": prompt}]
    )
    antwort = msg.content[0].text.strip()
    antwort = re.sub(r'^```(?:json)?\s*', '', antwort)
    antwort = re.sub(r'\s*```$', '', antwort).strip()
    try:
        data = json.loads(antwort)
        uebergaenge = data.get("uebergaenge", [])
        print(f"  Segmentierung via Claude ({len(uebergaenge)}/{len(agenda)} Punkte).")
        return uebergaenge
    except json.JSONDecodeError:
        print("  Segmentierung fehlgeschlagen – Fallback auf Einzel-Pass.")
        return []


# ── Word-Dokument erstellen ───────────────────────────────────
def erstelle_protokoll(transkript_pfad: str, thema: str,
                       agenda_pfad: str = None,
                       teilnehmer_liste: list = None,
                       extra_instruktionen: str = None,
                       kanal_wechsel: list = None,
                       teilnehmer_pro_kanal: dict = None):
    """
    Erstellt ein Word-Protokoll aus einem Transkript.

    Parameter:
        transkript_pfad:     Pfad zur Transkript-TXT-Datei
        thema:               Titel/Thema der Sitzung
        agenda_pfad:         Optionaler Pfad zur agenda.txt
        teilnehmer_liste:    Optionale Liste von Teilnehmern (Server-Modus).
                             Format: [{"name": "...", "frs": "..."}, ...]
                             Wenn übergeben, werden PNGs und Claude Vision übersprungen.
        extra_instruktionen: Freier Text, der zusätzlich an den Claude-Prompt angehängt wird.
    """
    skript_ordner  = os.path.dirname(os.path.abspath(__file__))
    config         = lese_config(skript_ordner)
    volltext, segmente = lese_transkript(transkript_pfad)
    agenda         = lese_agenda(agenda_pfad)
    api_key        = config.get("ANTHROPIC_API_KEY", "")
    modell         = config.get("CLAUDE_MODEL", "claude-sonnet-4-5-20250929")
    hat_api        = bool(api_key and not api_key.startswith("sk-ant-HIER"))

    # Ausgabepfad + Sitzungsdatum (wird für ki_zuordnung und Dokument benötigt)
    ausgabe_ordner = os.path.dirname(transkript_pfad)
    sitzungs_datum = datetime.now()
    meta_path = os.path.join(ausgabe_ordner, "meta.json")
    if os.path.exists(meta_path):
        try:
            with open(meta_path, encoding="utf-8") as f:
                _meta = json.load(f)
            if _meta.get("started_at"):
                sitzungs_datum = datetime.fromisoformat(_meta["started_at"])
        except Exception:
            pass

    # Teilnehmer-Ermittlung:
    # Server-Modus: Liste direkt übergeben → Vision entfällt
    # Windows-Modus: PNGs im Skript-Ordner suchen und per Vision auswerten
    if teilnehmer_liste is None:
        pngs = finde_alle_pngs(skript_ordner)
        if pngs and hat_api:
            teilnehmer_liste = extrahiere_teilnehmer(pngs, api_key, modell)
            print(f"{len(teilnehmer_liste)} Teilnehmer gefunden.")
        elif pngs:
            print("HINWEIS: Kein API-Key – Screenshots werden ohne OCR übersprungen.")
            teilnehmer_liste = []
        else:
            teilnehmer_liste = []
    # else: Liste aus Server-Query → direkt verwenden, kein Vision-Call nötig

    # KI-Zuordnung Transkript → Agenda
    ki_punkte = []
    if agenda and hat_api:
        # Alle Teilnehmer aus beiden Quellen zusammenführen
        alle_teilnehmer = list(teilnehmer_liste or [])
        if teilnehmer_pro_kanal:
            seen = {t.get("frs") or t.get("name") for t in alle_teilnehmer}
            for tl in teilnehmer_pro_kanal.values():
                for t in tl:
                    key = t.get("frs") or t.get("name")
                    if key not in seen:
                        alle_teilnehmer.append(t)
                        seen.add(key)

        ki_punkte = ki_zuordnung(
            volltext, segmente, agenda, api_key, modell,
            extra_instruktionen=extra_instruktionen,
            kanal_wechsel=kanal_wechsel or [],
            teilnehmer=alle_teilnehmer,
            session_started_at=sitzungs_datum,
        )
        if ki_punkte:
            print(f"Claude hat {len(ki_punkte)} Agenda-Punkte zugeordnet.")
    elif agenda:
        print("HINWEIS: Kein API-Key – Protokoll ohne KI-Zuordnung.")

    # Ausgabe-Datei – Datum aus Session-Ordnernamen ableiten, nicht aus datetime.now()
    session_folder_name = Path(ausgabe_ordner).name  # z.B. "20260301_200000"
    try:
        zeitstempel = session_folder_name[:13]  # "20260301_2000" = YYYYMMDD_HHMM
    except Exception:
        zeitstempel = datetime.now().strftime("%Y%m%d_%H%M")
    ausgabe_datei = os.path.join(ausgabe_ordner, f"Protokoll_{zeitstempel}.docx")

    doc = Document()
    setze_seitenraender(doc)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    # ── Titel ──────────────────────────────────────────────────
    titel = doc.add_heading('Sitzungsprotokoll', level=1)
    titel.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titel.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    if thema:
        ut = doc.add_heading(thema, level=2)
        ut.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ── Metadaten ──────────────────────────────────────────────
    t = doc.add_table(rows=2, cols=2)
    t.style = 'Table Grid'
    for i, (label, wert) in enumerate([
        ("Datum:",       sitzungs_datum.strftime("%d.%m.%Y")),
        ("Erstellt am:", datetime.now().strftime("%d.%m.%Y %H:%M Uhr")),
    ]):
        t.cell(i, 0).text = label
        t.cell(i, 1).text = wert
        t.cell(i, 0).paragraphs[0].runs[0].bold = True
    doc.add_paragraph()

    # ── Kanalwechsel-Hinweis ──────────────────────────────────
    if kanal_wechsel:
        note = doc.add_paragraph()
        note.add_run("Hinweis – Kanalwechsel während der Sitzung:").bold = True
        for evt in kanal_wechsel:
            try:
                ts = datetime.fromisoformat(evt["timestamp"]).strftime("%H:%M Uhr")
            except Exception:
                ts = evt.get("timestamp", "")[:19].replace("T", " ")
            von  = evt.get("from_channel_name") or f"Kanal {evt.get('from_channel', '?')}"
            nach = evt.get("to_channel_name")   or f"Kanal {evt.get('to_channel', '?')}"
            item = doc.add_paragraph(style="List Bullet")
            item.add_run(f"{ts}: {von} → {nach}")
        doc.add_paragraph()

    # ── Inhaltsverzeichnis ─────────────────────────────────────
    doc.add_heading('Inhaltsverzeichnis', level=1)
    from docx.oxml import OxmlElement
    toc_para = doc.add_paragraph()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), ' TOC \\o "1-2" \\h \\z \\u ')
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    run.append(rPr)
    t = OxmlElement('w:t')
    t.text = '[Inhaltsverzeichnis – bitte in Word mit F9 aktualisieren]'
    run.append(t)
    fld.append(run)
    toc_para._p.append(fld)
    doc.add_paragraph()

    # ── Teilnehmer ─────────────────────────────────────────────
    doc.add_heading('Teilnehmer', level=1)

    def _teilnehmer_tabelle(dok, liste):
        tbl = dok.add_table(rows=1, cols=2)
        tbl.style = 'Table Grid'
        hdr = tbl.rows[0].cells
        hdr[0].text = "Name"
        hdr[1].text = "Callsign"
        for r in hdr:
            r.paragraphs[0].runs[0].bold = True
        for t in liste:
            row = tbl.add_row().cells
            row[0].text = t.get("name", "")
            row[1].text = t.get("frs", "")

    if teilnehmer_pro_kanal:
        for kanal_name, liste in teilnehmer_pro_kanal.items():
            if not liste:
                continue
            doc.add_heading(kanal_name, level=2)
            _teilnehmer_tabelle(doc, liste)
            doc.add_paragraph()
    elif teilnehmer_liste:
        _teilnehmer_tabelle(doc, teilnehmer_liste)
        doc.add_paragraph()
    else:
        for _ in range(4):
            doc.add_paragraph().add_run("_" * 50)
        doc.add_paragraph()

    # ── Agenda ─────────────────────────────────────────────────
    if agenda:
        doc.add_heading('Agenda', level=1)
        for i, punkt in enumerate(agenda, 1):
            p = doc.add_paragraph()
            p.add_run(f"{i}. {punkt}")
        doc.add_paragraph()

    # ── Protokoll ─────────────────────────────────────────────
    doc.add_heading('Protokoll', level=1)

    if ki_punkte:
        for eintrag in ki_punkte:
            h = doc.add_heading(eintrag.get("punkt", ""), level=2)
            h.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

            if eintrag.get("zeitraum"):
                zeitraum_str = _mm_ss_zu_uhrzeit(eintrag["zeitraum"], sitzungs_datum)
                p = doc.add_paragraph()
                r = p.add_run(f"Zeitraum: {zeitraum_str}")
                r.italic = True
                r.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
                r.font.size = Pt(9)

            if eintrag.get("zusammenfassung"):
                doc.add_paragraph(eintrag["zusammenfassung"])

            if eintrag.get("details"):
                for d in eintrag["details"]:
                    d_clean = re.sub(r'^[•\-–—]\s*', '', d.strip())
                    if d_clean.rstrip().endswith(':'):
                        # Unterüberschrift (Heading 3 – nicht im TOC wegen "1-2"-Begrenzung)
                        h3 = doc.add_heading(d_clean, level=3)
                        h3.runs[0].font.color.rgb = RGBColor(0x2D, 0x3A, 0x4A)
                    else:
                        doc.add_paragraph(d_clean, style='List Bullet')

            if eintrag.get("beschluesse"):
                h3_b = doc.add_heading("Beschlüsse / Aktionspunkte:", level=3)
                h3_b.runs[0].font.color.rgb = RGBColor(0x2D, 0x3A, 0x4A)
                for b in eintrag["beschluesse"]:
                    b_clean = re.sub(r'^[•\-–—]\s*', '', b.strip())
                    doc.add_paragraph(b_clean, style='List Bullet')

            doc.add_paragraph()

    elif agenda:
        for punkt in agenda:
            h = doc.add_heading(punkt, level=2)
            h.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
            for _ in range(3):
                doc.add_paragraph().add_run("_" * 70)
            doc.add_paragraph()
    else:
        for _ in range(6):
            doc.add_paragraph().add_run("_" * 70)

    doc.save(ausgabe_datei)
    print(f"\nProtokoll gespeichert: {ausgabe_datei}")

    # Alte Protokoll-Dateien mit abweichendem Namen entfernen (z.B. nach Umbenennung durch Regen)
    for alte_datei in Path(ausgabe_ordner).glob("Protokoll_*.docx"):
        if alte_datei != Path(ausgabe_datei):
            alte_datei.unlink()
            print(f"Altes Protokoll entfernt: {alte_datei.name}")

    return ausgabe_datei


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Verwendung: python protokoll_erstellen.py <transkript.txt> [thema] [agenda.txt]")
        sys.exit(1)

    erstelle_protokoll(
        transkript_pfad = sys.argv[1],
        thema           = sys.argv[2] if len(sys.argv) > 2 else "",
        agenda_pfad     = sys.argv[3] if len(sys.argv) > 3 else None,
    )
